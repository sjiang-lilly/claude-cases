#!/usr/bin/env python3
"""
Generate HER2 Epitope Analysis Report in Word format.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def create_report():
    """Generate comprehensive Word report."""

    doc = Document()

    # ========== TITLE PAGE ==========
    title = doc.add_heading('HER2 Epitope Analysis for ADC Binder Design', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    subtitle = doc.add_paragraph('Comprehensive Analysis of HER2 Epitopes, mAbs, and Resistance Mechanisms')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Author info
    author_table = doc.add_table(rows=4, cols=2)
    author_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    author_data = [
        ('Author:', 'Mandy Jiang'),
        ('Email:', 'shan.jiang2@lilly.com'),
        ('Affiliation:', 'Eli Lilly and Company - Oncology, Bioinformatics'),
        ('Date:', datetime.now().strftime('%Y-%m-%d'))
    ]
    for i, (label, value) in enumerate(author_data):
        author_table.rows[i].cells[0].text = label
        author_table.rows[i].cells[1].text = value

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. HER2 Biology and Domain Structure',
        '3. HER2 Epitope Analysis',
        '4. Approved mAbs and ADCs',
        '5. Internalization Predictions',
        '6. Mutation and Resistance Analysis',
        '7. Scientific Plan for Overcoming Resistance',
        '8. Recommendations',
        '9. Methods',
        '10. References'
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ========== 1. EXECUTIVE SUMMARY ==========
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_paragraph(
        'This report presents a comprehensive analysis of HER2 epitopes for antibody-drug conjugate '
        '(ADC) binder design. HER2 (ErbB2/ERBB2) is a validated oncology target overexpressed in '
        '~20% of breast cancers and various other malignancies. The analysis covers epitope mapping, '
        'approved therapeutics, internalization mechanisms, and resistance strategies.'
    )

    doc.add_heading('Key Findings', level=2)
    findings = [
        'Domain IV epitope (residues 557-603) is the primary ADC target, used by T-DM1 and T-DXd',
        'Biparatopic antibodies (targeting Domain II + IV) show superior internalization (70% vs 25%)',
        'HER2 downregulation is the major resistance mechanism (15-30%), not epitope mutations',
        'T-DXd bystander effect enables efficacy in HER2-low tumors',
        'Next-generation biparatopic ADCs (ZW49) offer promising solutions for resistance'
    ]
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')

    doc.add_page_break()

    # ========== 2. HER2 BIOLOGY ==========
    doc.add_heading('2. HER2 Biology and Domain Structure', level=1)

    doc.add_paragraph(
        'HER2 (Human Epidermal Growth Factor Receptor 2) is a 185 kDa transmembrane receptor '
        'tyrosine kinase belonging to the ErbB family. Unlike other family members, HER2 has '
        'no known ligand and functions primarily as a co-receptor through heterodimerization.'
    )

    doc.add_heading('Domain Structure', level=2)

    # Domain table
    domain_table = doc.add_table(rows=5, cols=4)
    domain_table.style = 'Table Grid'
    headers = ['Domain', 'Residues', 'Function', 'Therapeutic Relevance']
    for i, header in enumerate(headers):
        domain_table.rows[0].cells[i].text = header

    domain_data = [
        ('Domain I', '23-195', 'L1 domain, dimerization', 'Zanidatamab target'),
        ('Domain II', '196-319', 'Cysteine-rich, dimerization arm', 'Pertuzumab epitope'),
        ('Domain III', '320-488', 'L2 domain, ligand binding', 'Experimental targets'),
        ('Domain IV', '489-630', 'Membrane proximal', 'Trastuzumab/ADC epitope'),
    ]
    for i, row_data in enumerate(domain_data):
        for j, cell_data in enumerate(row_data):
            domain_table.rows[i+1].cells[j].text = cell_data

    # Add schematic image
    doc.add_paragraph()
    if os.path.exists('images/her2_domain_schematic.png'):
        doc.add_picture('images/her2_domain_schematic.png', width=Inches(6))
        doc.add_paragraph('Figure 1: HER2 domain structure and therapeutic antibody binding sites',
                         style='Caption')

    doc.add_page_break()

    # ========== 3. EPITOPE ANALYSIS ==========
    doc.add_heading('3. HER2 Epitope Analysis', level=1)

    doc.add_paragraph(
        'Five major epitope regions have been characterized on HER2, each with distinct '
        'properties relevant to ADC design.'
    )

    # Epitope table
    epitope_table = doc.add_table(rows=6, cols=5)
    epitope_table.style = 'Table Grid'
    headers = ['Epitope', 'Domain', 'Residues', 'mAbs', 'ADC Suitability']
    for i, header in enumerate(headers):
        epitope_table.rows[0].cells[i].text = header

    epitope_data = [
        ('EPI-001', 'Domain IV', '557-603', 'Trastuzumab, T-DM1, T-DXd', '8.8/10'),
        ('EPI-002', 'Domain II', '266-333', 'Pertuzumab', '7.8/10'),
        ('EPI-003', 'Domain I', '23-165', 'Zanidatamab', '7.8/10'),
        ('EPI-004', 'Domain III', '355-435', 'Experimental', '6.0/10'),
        ('EPI-005', 'Biparatopic', 'II + IV', 'Zanidatamab', '9.5/10'),
    ]
    for i, row_data in enumerate(epitope_data):
        for j, cell_data in enumerate(row_data):
            epitope_table.rows[i+1].cells[j].text = cell_data

    doc.add_paragraph()
    doc.add_heading('Epitope Evaluation Criteria', level=2)
    criteria = [
        'Accessibility: Surface exposure for antibody binding',
        'Internalization: Rate of receptor-mediated endocytosis',
        'Clinical validation: Evidence from approved therapeutics',
        'Stability: Conformational stability of epitope region'
    ]
    for c in criteria:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_page_break()

    # ========== 4. APPROVED mAbs AND ADCs ==========
    doc.add_heading('4. Approved mAbs and ADCs', level=1)

    doc.add_heading('4.1 Monoclonal Antibodies', level=2)

    mab_table = doc.add_table(rows=4, cols=5)
    mab_table.style = 'Table Grid'
    headers = ['mAb', 'Epitope', 'Kd (nM)', 'Mechanism', 'Approval']
    for i, header in enumerate(headers):
        mab_table.rows[0].cells[i].text = header

    mab_data = [
        ('Trastuzumab', 'Domain IV', '5.0', 'ADCC, signaling block', 'FDA 1998'),
        ('Pertuzumab', 'Domain II', '1.0', 'Blocks dimerization', 'FDA 2012'),
        ('Margetuximab', 'Domain IV', '4.8', 'Enhanced ADCC (Fc opt)', 'FDA 2020'),
    ]
    for i, row_data in enumerate(mab_data):
        for j, cell_data in enumerate(row_data):
            mab_table.rows[i+1].cells[j].text = cell_data

    doc.add_paragraph()
    doc.add_heading('4.2 Antibody-Drug Conjugates', level=2)

    adc_table = doc.add_table(rows=4, cols=6)
    adc_table.style = 'Table Grid'
    headers = ['ADC', 'Linker', 'Payload', 'DAR', 'Approval', 'Indication']
    for i, header in enumerate(headers):
        adc_table.rows[0].cells[i].text = header

    adc_data = [
        ('T-DM1 (Kadcyla)', 'Non-cleavable', 'DM1', '3.5', 'FDA 2013', 'HER2+ mBC'),
        ('T-DXd (Enhertu)', 'Cleavable', 'DXd', '8.0', 'FDA 2019', 'HER2+/low'),
        ('Disitamab ved.', 'Cleavable', 'MMAE', '4.0', 'China 2021', 'HER2+ GC'),
    ]
    for i, row_data in enumerate(adc_data):
        for j, cell_data in enumerate(row_data):
            adc_table.rows[i+1].cells[j].text = cell_data

    # Add mAb summary image
    doc.add_paragraph()
    if os.path.exists('images/mab_summary_table.png'):
        doc.add_picture('images/mab_summary_table.png', width=Inches(6))
        doc.add_paragraph('Figure 2: Summary of HER2-targeting mAbs and ADCs', style='Caption')

    doc.add_page_break()

    # ========== 5. INTERNALIZATION ==========
    doc.add_heading('5. Internalization Predictions', level=1)

    doc.add_paragraph(
        'Receptor internalization is critical for ADC efficacy. The rate of endocytosis '
        'and subsequent lysosomal delivery determines payload release efficiency.'
    )

    intern_table = doc.add_table(rows=4, cols=5)
    intern_table.style = 'Table Grid'
    headers = ['Epitope', 'Rate', '4h Uptake', 'Recycling', 'ADC Score']
    for i, header in enumerate(headers):
        intern_table.rows[0].cells[i].text = header

    intern_data = [
        ('Domain IV', 'Slow', '25%', '60-70%', '7.5/10'),
        ('Domain II', 'Very slow', '15%', '80%', '5.0/10'),
        ('Biparatopic', 'Fast', '70%', '20%', '9.5/10'),
    ]
    for i, row_data in enumerate(intern_data):
        for j, cell_data in enumerate(row_data):
            intern_table.rows[i+1].cells[j].text = cell_data

    # Add comparison chart
    doc.add_paragraph()
    if os.path.exists('images/epitope_comparison.png'):
        doc.add_picture('images/epitope_comparison.png', width=Inches(6))
        doc.add_paragraph('Figure 3: Epitope-dependent internalization and ADC suitability',
                         style='Caption')

    doc.add_heading('Key Insight', level=2)
    doc.add_paragraph(
        'Biparatopic antibodies (like Zanidatamab) induce receptor clustering, dramatically '
        'enhancing internalization from 25% to 70% at 4 hours. This makes biparatopic ADCs '
        'ideal candidates for next-generation HER2-targeted therapy.'
    )

    doc.add_page_break()

    # ========== 6. RESISTANCE ANALYSIS ==========
    doc.add_heading('6. Mutation and Resistance Analysis', level=1)

    doc.add_heading('6.1 Point Mutations', level=2)
    doc.add_paragraph(
        'Most HER2 point mutations occur in the kinase domain and do NOT affect ADC binding. '
        'These mutations primarily confer TKI resistance while leaving ADC efficacy intact.'
    )

    mut_table = doc.add_table(rows=5, cols=4)
    mut_table.style = 'Table Grid'
    headers = ['Mutation', 'Domain', 'Frequency', 'ADC Impact']
    for i, header in enumerate(headers):
        mut_table.rows[0].cells[i].text = header

    mut_data = [
        ('L755S', 'Kinase', '2.1%', 'None'),
        ('V777L', 'Kinase', '1.5%', 'None'),
        ('S310F', 'Domain II', '0.8%', 'Minimal'),
        ('T798M', 'Kinase', '0.5%', 'None'),
    ]
    for i, row_data in enumerate(mut_data):
        for j, cell_data in enumerate(row_data):
            mut_table.rows[i+1].cells[j].text = cell_data

    doc.add_heading('6.2 Major Resistance Mechanisms', level=2)
    resist_items = [
        'HER2 downregulation (15-30%): Loss of target expression',
        'p95-HER2 truncation (20-30%): Loss of extracellular domain',
        'MDR1/P-gp efflux (T-DM1): Drug efflux from cells',
        'Tumor heterogeneity: Mixed HER2 expression levels'
    ]
    for item in resist_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========== 7. SCIENTIFIC PLAN ==========
    doc.add_heading('7. Scientific Plan for Overcoming Resistance', level=1)

    strategies = [
        ('Biparatopic ADCs', 'Target multiple epitopes (Domain II + IV) to reduce escape and enhance internalization'),
        ('Bystander Effect', 'Use membrane-permeable payloads (T-DXd) to kill HER2-negative bystander cells'),
        ('Combination Therapy', 'Combine ADC with checkpoint inhibitors or TKIs'),
        ('Novel Epitopes', 'Develop antibodies targeting underexplored Domain I or III epitopes'),
        ('Alternative Targets', 'For p95-HER2: target TROP2, HER3, or use CAR-T approaches')
    ]

    for title, desc in strategies:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ========== 8. RECOMMENDATIONS ==========
    doc.add_heading('8. Recommendations', level=1)

    doc.add_heading('For ADC Development', level=2)
    dev_recs = [
        'Prioritize biparatopic antibody platforms for new ADC development',
        'Use cleavable linkers with membrane-permeable payloads',
        'Target high DAR (≥8) to compensate for slow internalization',
        'Develop companion diagnostics for HER2 dynamics monitoring'
    ]
    for rec in dev_recs:
        doc.add_paragraph(rec, style='List Bullet')

    doc.add_heading('For Clinical Application', level=2)
    clin_recs = [
        'Screen for p95-HER2 status before initiating ADC therapy',
        'Monitor HER2 expression longitudinally with liquid biopsy',
        'Consider T-DXd for HER2-low patients (bystander effect)',
        'Sequence ADCs after TKI failure (different resistance mechanisms)'
    ]
    for rec in clin_recs:
        doc.add_paragraph(rec, style='List Bullet')

    doc.add_page_break()

    # ========== 9. METHODS ==========
    doc.add_heading('9. Methods', level=1)

    doc.add_heading('Data Sources', level=2)
    doc.add_paragraph(
        'HER2 protein sequence was retrieved from UniProt (P04626). Crystal structures '
        'were obtained from RCSB PDB (1N8Z, 1S78, 6OGE). Mutation data was compiled from '
        'COSMIC and published literature.'
    )

    doc.add_heading('Epitope Analysis', level=2)
    doc.add_paragraph(
        'Epitope regions were mapped based on published crystallographic studies of '
        'antibody-HER2 complexes. ADC suitability scores were calculated using weighted '
        'criteria including accessibility, internalization rate, clinical validation, '
        'and structural stability.'
    )

    doc.add_heading('Internalization Predictions', level=2)
    doc.add_paragraph(
        'Internalization rates were compiled from published flow cytometry and imaging '
        'studies. Predictions for novel epitopes were based on structural homology and '
        'receptor trafficking literature.'
    )

    doc.add_heading('Tools and Dependencies', level=2)
    tools = [
        'Python 3.x with BioPython, pandas, numpy',
        'py3Dmol for interactive 3D visualization',
        'matplotlib/seaborn for static figures',
        'python-docx for report generation'
    ]
    for tool in tools:
        doc.add_paragraph(tool, style='List Bullet')

    doc.add_page_break()

    # ========== 10. REFERENCES ==========
    doc.add_heading('10. References', level=1)

    references = [
        '1. Cho HS, et al. Structure of the extracellular region of HER2 alone and in complex with the Herceptin Fab. Nature. 2003;421:756-760.',
        '2. Franklin MC, et al. Insights into ErbB signaling from the structure of the ErbB2-pertuzumab complex. Cancer Cell. 2004;5:317-328.',
        '3. Lewis Phillips GD, et al. Targeting HER2-positive breast cancer with trastuzumab-DM1. Cancer Res. 2008;68:9280-9290.',
        '4. Modi S, et al. Trastuzumab deruxtecan in previously treated HER2-positive breast cancer. N Engl J Med. 2020;382:610-621.',
        '5. Modi S, et al. Trastuzumab deruxtecan in previously treated HER2-low advanced breast cancer. N Engl J Med. 2022;387:9-20.',
        '6. Weisser NE, et al. An anti-HER2 biparatopic antibody that induces unique HER2 clustering. Nat Commun. 2023;14:1394.',
        '7. Hunter FW, et al. Mechanisms of resistance to antibody-drug conjugates. Cancer Res. 2020;80:5057-5067.',
        '8. Scaltriti M, et al. Expression of p95HER2, a truncated form of the HER2 receptor. J Natl Cancer Inst. 2007;99:628-638.',
    ]

    for ref in references:
        doc.add_paragraph(ref, style='List Number')

    # Save document
    output_path = 'output/HER2_Epitope_Report.docx'
    doc.save(output_path)
    print(f"Generated: {output_path}")

    return output_path


if __name__ == "__main__":
    create_report()
