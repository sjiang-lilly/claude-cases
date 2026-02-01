#!/usr/bin/env python3
"""
Add 2025 Hu et al. Nature Cancer reference to HER2 Epitope Report.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_2025_reference(doc):
    """Add the 2025 Hu et al. reference and findings to the document."""

    # Add page break
    doc.add_page_break()

    # Section header
    heading = doc.add_heading('p95-HER2 Immunosuppressive Program (2025 Update)', level=1)
    heading.runs[0].font.color.rgb = RGBColor(192, 0, 0)

    # Key reference
    doc.add_heading('Critical New Finding', level=2)

    ref_para = doc.add_paragraph()
    ref_para.add_run('Hu D, et al. ').bold = True
    ref_para.add_run(
        'p95HER2, a truncated form of the HER2 oncoprotein, drives an immunosuppressive '
        'program in HER2+ breast cancer that limits trastuzumab deruxtecan efficacy. '
    ).italic = True
    ref_para.add_run('Nat Cancer. 2025. doi:10.1038/s43018-025-00969-4. PMID: 40579589')

    doc.add_paragraph()

    # Key findings table
    doc.add_heading('Key Findings', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    # Header row
    headers = ['Mechanism', 'Effect', 'Implication']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Data rows
    data = [
        ('PD-L1 upregulation', 'Immune checkpoint activation', 'T cell exhaustion'),
        ('IL-6 secretion', 'Immunosuppressive cytokine', 'Reduced anti-tumor immunity'),
        ('T-DXd resistance', 'Limited bystander effect', 'Reduced ADC efficacy'),
        ('Neratinib sensitivity', 'Degrades p95HER2', 'Potential combination therapy'),
    ]

    for i, (mechanism, effect, implication) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = mechanism
        row.cells[1].text = effect
        row.cells[2].text = implication

    doc.add_paragraph()

    # Therapeutic implications
    doc.add_heading('Therapeutic Implications', level=2)

    implications = [
        'p95HER2 creates an immunosuppressive tumor microenvironment',
        'This directly limits T-DXd (trastuzumab deruxtecan) efficacy',
        'Neratinib can degrade p95HER2, restoring immune response',
        'Combination of p95-targeting ADC + immune checkpoint inhibitors may be synergistic',
        'Strongly supports rationale for p95-Bispecific-001 development',
    ]

    for imp in implications:
        doc.add_paragraph(imp, style='List Bullet')

    doc.add_paragraph()

    # Highlight box
    highlight = doc.add_paragraph()
    highlight.add_run('SIGNIFICANCE: ').bold = True
    highlight.add_run(
        'This 2025 Nature Cancer paper provides mechanistic evidence that p95-HER2 '
        'is not just a passive resistance marker, but actively drives immunosuppression '
        'that undermines current ADC therapies. This strengthens the case for developing '
        'p95-specific ADCs like our predicted p95-Bispecific-001.'
    )

    return doc


def main():
    """Update the HER2 Epitope Report with 2025 reference."""

    input_path = "../output/HER2_Epitope_Report.docx"
    output_path = "../output/HER2_Epitope_Report.docx"

    print("Adding 2025 Hu et al. Nature Cancer reference to report...")

    if os.path.exists(input_path):
        doc = Document(input_path)
    else:
        print(f"Warning: {input_path} not found. Creating new document.")
        doc = Document()
        doc.add_heading('HER2 Epitope Analysis for ADC Binder Design', level=0)

    # Add the 2025 reference section
    doc = add_2025_reference(doc)

    # Save
    doc.save(output_path)
    print(f"Saved: {output_path}")

    print("\nNew section added:")
    print("  - p95-HER2 Immunosuppressive Program (2025 Update)")
    print("  - Hu et al. Nature Cancer 2025 reference")
    print("  - Key findings table")
    print("  - Therapeutic implications")


if __name__ == "__main__":
    main()
